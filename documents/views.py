import os
from datetime import datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from .forms import UploadFileForm
from .models import Document
from docx import Document as DocxDocument
from fpdf import FPDF
from pymediainfo import MediaInfo


def upload_file(request):
    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES["file"]
            document = Document(file=uploaded_file)
            document.save()  # Save the document instance to the database

            file_path = document.file.path

            basic_metadata = extract_basic_metadata(file_path)
            detailed_metadata = extract_detailed_metadata(file_path)
            file_metadata = {**basic_metadata, **detailed_metadata}

            try:
                convert_to_pdf(document)
            except Exception as e:
                return HttpResponse(f"Error converting file to PDF: {e}", status=500)

            return render(request, "success.html", {"file_metadata": file_metadata, "document": document})
    else:
        form = UploadFileForm()
    return render(request, "upload.html", {"form": form})


def success(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    return render(request, 'success.html', {'document': document})


def convert_to_pdf(document):
    docx_path = document.file.path
    pdf_path = docx_path.replace('.docx', '.pdf')

    try:
        doc = DocxDocument(docx_path)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=12)

        for paragraph in doc.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)

        pdf.output(pdf_path)

        document.pdf_file.name = os.path.relpath(pdf_path, start='media')  # Save relative path
        document.save()
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        raise e


def download_pdf(request, doc_id):
    document = get_object_or_404(Document, id=doc_id)
    
    if not document.pdf_file or not document.pdf_file.path:
        return HttpResponse("The PDF file does not exist.", status=404)

    try:
        with open(document.pdf_file.path, 'rb') as f:
            response = HttpResponse(f, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{os.path.basename(document.pdf_file.path)}"'
            return response
    except FileNotFoundError:
        return HttpResponse("The PDF file could not be found.", status=404)


def extract_basic_metadata(file_path):
    return {
        "name": os.path.basename(file_path),
        "size": os.path.getsize(file_path),  # Size in bytes
        "creation_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime("%Y-%m-%d %H:%M:%S"),
        "modification_time": datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S"),
        "path": file_path,
    }


def extract_detailed_metadata(file_path):
    try:
        media_info = MediaInfo.parse(file_path)
        metadata = {}
        for track in media_info.tracks:
            if track.track_type == "General":
                metadata["duration"] = f"{int(track.duration / 1000)} seconds" if track.duration else "N/A"
                metadata["file_size"] = f"{int(track.file_size / 1024)} KB" if track.file_size else "N/A"
                metadata["file_type"] = track.format if track.format else "Unknown"
        return metadata
    except Exception as e:
        print(f"Error extracting detailed metadata: {e}")
        return {
            "duration": "N/A",
            "file_size": "N/A",
            "file_type": "N/A",
        }
